# -*- coding: utf-8 -*-
"""Template de site de cours — Mohamed EL AFRIT (MIT / CC BY 4.0)

Générateur de polycopiés de synthèse imprimables (DOCX + PDF), bilingue FR/EN,
pour le TEMPLATE de site de cours. Identité visuelle : bleu nuit / doré.

Le contenu est lu depuis ``polycopies_content.json`` (même dossier) : une source
commune bilingue produit des rendus cohérents dans les deux langues.

Usage :
    python generate_polycopies.py            # FR + EN (défaut)
    python generate_polycopies.py fr         # FR seulement
    python generate_polycopies.py en         # EN seulement
    python generate_polycopies.py both       # FR + EN

Sorties : tools/out/chapitre-<n>-polycopie.<lang>.docx et .pdf

Dépendances : python-docx, fpdf2 (+ pypdfium2 facultatif).
"""
import os
import re
import sys
import json

from fpdf import FPDF
from fpdf.fonts import FontFace
import docx
from docx.shared import Pt, RGBColor, Mm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
OUT = os.path.join(HERE, 'out')
CONTENT = os.path.join(HERE, 'polycopies_content.json')
FONTS = r'C:\Windows\Fonts'

# ---------------- Identité visuelle (bleu nuit / doré) ----------------
NAVY = (27, 58, 91); NAVY_D = (14, 34, 56); GOLD = (194, 161, 77); GOLD_D = (156, 127, 51)
INK = (22, 32, 46); SLATE = (60, 74, 92)
CALLOUT = {  # type -> (bordure RVB, fond clair RVB, hex_docx_fond)
  'info':    ((47, 111, 159), (233, 240, 247), 'E9F0F7'),
  'tip':     ((47, 143, 91),  (233, 245, 238), 'E9F5EE'),
  'warning': ((201, 138, 43), (250, 243, 230), 'FAF3E6'),
  'danger':  ((192, 57, 43),  (250, 235, 233), 'FAEBE9'),
  'legal':   ((107, 79, 160), (240, 236, 247), 'F0ECF7'),
}

# ---------------- Localisation des libellés d'interface ----------------
LANGS = ('fr', 'en')
LABELS = {
  'fr': {
    'kicker':    "POLYCOPIÉ DE SYNTHÈSE — CHAPITRE %d",
    'subtitle':  "Polycopié de synthèse",
    'toc':       "Sommaire",
    'takeaways': "À retenir",
    'lawref':    "Référence",
    'footer':    "Mohamed EL AFRIT — CC BY 4.0",
    'page':      "Page %d",
  },
  'en': {
    'kicker':    "SUMMARY HANDOUT — CHAPTER %d",
    'subtitle':  "Summary handout",
    'toc':       "Contents",
    'takeaways': "Key takeaways",
    'lawref':    "Reference",
    'footer':    "Mohamed EL AFRIT — CC BY 4.0",
    'page':      "Page %d",
  },
}


# ---------------- Sélection de langue & accès au contenu ----------------
def pick(value, lang):
    """Renvoie la variante de langue d'un champ bilingue.

    Accepte un dict {"fr":..., "en":...} ou une simple chaîne (repli commun).
    Repli : langue demandée -> 'fr' -> 'en' -> première valeur disponible.
    """
    if isinstance(value, dict):
        if lang in value and value[lang] is not None:
            return value[lang]
        for fb in ('fr', 'en'):
            if value.get(fb):
                return value[fb]
        return next(iter(value.values()), '')
    return value if value is not None else ''


def load_chapters():
    if not os.path.exists(CONTENT):
        raise SystemExit("Fichier de contenu introuvable : %s" % CONTENT)
    with open(CONTENT, encoding='utf-8') as f:
        raw = json.load(f)
    # Tri numérique des clés de chapitre ("1", "2", ...)
    return {int(k): v for k, v in sorted(raw.items(), key=lambda kv: int(kv[0]))}


def toc(blocks, lang):
    return [pick(b[1], lang) for b in blocks if b[0] == 'h2']


def split_bold(text):
    """Découpe un texte balisé **gras** -> liste de (segment, gras?)."""
    out = []; bold = False
    for part in (text or '').split('**'):
        if part:
            out.append((part, bold))
        bold = not bold
    return out


def ref_label(ref):
    """Retire un préfixe « Référence/Reference — » déjà présent (évite le doublon)."""
    return re.sub(r'^\s*(r[ée]f[ée]rence|reference)\s*[—\-:]\s*', '', ref or '', flags=re.I).strip()


# ======================= PDF (fpdf2) =======================
class Poly(FPDF):
    def header(self):
        if self.page_no() == 1:
            return
        self.set_font('Ar', '', 8); self.set_text_color(*SLATE)
        self.cell(0, 6, self.title_short, align='L')
        self.ln(7)

    def footer(self):
        self.set_y(-14); self.set_font('Ar', '', 8); self.set_text_color(*SLATE)
        self.cell(0, 6, self.foot_left, align='L')
        self.cell(0, 6, self.page_label % self.page_no(), align='R')


def _register_fonts(pdf):
    """Enregistre Arial (Unicode) avec repli sur arial.ttf si une variante manque."""
    base = os.path.join(FONTS, 'arial.ttf')
    for st, fn in [('', 'arial.ttf'), ('B', 'arialbd.ttf'), ('I', 'ariali.ttf'), ('BI', 'arialbi.ttf')]:
        p = os.path.join(FONTS, fn)
        pdf.add_font('Ar', st, p if os.path.exists(p) else base)


def pdf_rich(pdf, text, size=10.5, lh=5.6):
    pdf.set_text_color(*INK)
    for seg, bold in split_bold(text):
        pdf.set_font('Ar', 'B' if bold else '', size)
        pdf.write(lh, seg)
    pdf.ln(lh + 2)


def mc_height(pdf, text, w, lh, style='', size=10, markdown=False):
    """Hauteur d'un multi_cell sans rien dessiner (évite qu'un encadré soit coupé)."""
    pdf.set_font('Ar', style, size)
    return len(pdf.multi_cell(w, lh, text, dry_run=True, output="LINES", markdown=markdown)) * lh


def render_pdf(num, ch, lang, path):
    L = LABELS[lang]
    title = pick(ch['title'], lang)
    subtitle = pick(ch.get('subtitle'), lang) or L['subtitle']

    pdf = Poly(format='A4')
    pdf.title_short = "%s %d — %s" % ("Chapitre" if lang == 'fr' else "Chapter", num, title)
    pdf.foot_left = L['footer']
    pdf.page_label = L['page']
    pdf.set_margins(18, 16, 18); pdf.set_auto_page_break(True, 18)
    _register_fonts(pdf)
    pdf.set_lang(lang)
    pdf.add_page()
    W = pdf.epw  # largeur utile

    # --- Bandeau titre ---
    pdf.set_fill_color(*NAVY); pdf.rect(0, 0, 210, 52, 'F')
    pdf.set_fill_color(*GOLD); pdf.rect(0, 52, 210, 1.5, 'F')
    pdf.set_xy(18, 14); pdf.set_text_color(*GOLD); pdf.set_font('Ar', 'B', 10)
    pdf.cell(0, 6, L['kicker'] % num)
    pdf.set_xy(18, 22); pdf.set_text_color(255, 255, 255); pdf.set_font('Ar', 'B', 19)
    pdf.multi_cell(W, 8.5, title)
    pdf.set_x(18); pdf.set_font('Ar', '', 10.5); pdf.set_text_color(225, 230, 238)
    pdf.multi_cell(W, 6, subtitle)
    pdf.set_xy(0, 58)

    # --- Mini-sommaire ---
    pdf.set_x(18); pdf.set_fill_color(244, 246, 249); pdf.set_draw_color(*GOLD)
    y0 = pdf.get_y()
    pdf.set_font('Ar', 'B', 10.5); pdf.set_text_color(*NAVY)
    pdf.multi_cell(W, 7, "  " + L['toc'], fill=True)
    pdf.set_font('Ar', '', 10); pdf.set_text_color(*SLATE)
    for t in toc(ch['blocks'], lang):
        pdf.set_x(18); pdf.multi_cell(W, 5.6, "   • " + t, fill=True)
    y1 = pdf.get_y()
    pdf.set_line_width(0.6); pdf.set_draw_color(*GOLD); pdf.line(18, y0, 18, y1)
    pdf.ln(5)

    # --- Blocs ---
    for b in ch['blocks']:
        kind = b[0]
        if kind == 'h2':
            heading = pick(b[1], lang)
            pdf.ln(2); pdf.start_section(heading)  # signet / outline cliquable
            pdf.set_x(18); pdf.set_font('Ar', 'B', 13.5); pdf.set_text_color(*NAVY)
            pdf.multi_cell(W, 7, heading)
            yy = pdf.get_y(); pdf.set_draw_color(*GOLD); pdf.set_line_width(0.5)
            pdf.line(18, yy + 0.5, 18 + 28, yy + 0.5)
            pdf.ln(3)
        elif kind == 'p':
            pdf.set_x(18); pdf_rich(pdf, pick(b[1], lang))
        elif kind == 'callout':
            _, ctype, ctitle, ctext = b
            ctitle = pick(ctitle, lang); ctext = pick(ctext, lang)
            bd, bg, _hx = CALLOUT.get(ctype, CALLOUT['info'])
            th = (mc_height(pdf, ctitle, W - 3, 6.2, 'B', 10.5)
                  + mc_height(pdf, ctext, W - 3, 5.4, '', 10, markdown=True) + 3)
            if pdf.will_page_break(th):
                pdf.add_page()
            pdf.ln(1); x0 = 18; y0 = pdf.get_y()
            pdf.set_fill_color(*bg); pdf.set_x(x0 + 3)
            pdf.set_font('Ar', 'B', 10.5); pdf.set_text_color(*bd)
            pdf.multi_cell(W - 3, 6.2, ctitle.replace('**', ''), fill=True)
            pdf.set_x(x0 + 3); pdf.set_font('Ar', '', 10); pdf.set_text_color(*INK)
            pdf.multi_cell(W - 3, 5.4, ctext, fill=True, markdown=True)
            y1 = pdf.get_y()
            pdf.set_fill_color(*bd); pdf.rect(x0, y0, 2.2, y1 - y0, 'F')
            pdf.ln(4)
        elif kind == 'lawref':
            _, ref, txt = b
            ref = pick(ref, lang); txt = pick(txt, lang)
            bd, bg, _hx = CALLOUT['legal']
            head = (L['lawref'] + " — " + ref_label(ref)).upper()
            th = (mc_height(pdf, head, W - 3, 6, 'B', 9.5)
                  + mc_height(pdf, "« " + txt + " »", W - 3, 5.4, 'I', 10, markdown=True) + 3)
            if pdf.will_page_break(th):
                pdf.add_page()
            pdf.ln(1); x0 = 18; y0 = pdf.get_y()
            pdf.set_fill_color(*bg); pdf.set_x(x0 + 3)
            pdf.set_font('Ar', 'B', 9.5); pdf.set_text_color(*bd)
            pdf.multi_cell(W - 3, 6, head, fill=True)
            pdf.set_x(x0 + 3); pdf.set_font('Ar', 'I', 10); pdf.set_text_color(*SLATE)
            pdf.multi_cell(W - 3, 5.4, "« " + txt + " »", fill=True, markdown=True)
            y1 = pdf.get_y(); pdf.set_fill_color(*bd); pdf.rect(x0, y0, 2.2, y1 - y0, 'F'); pdf.ln(4)
        elif kind == 'table':
            headers = [pick(h, lang) for h in b[1]]
            rows = [[pick(c, lang) for c in row] for row in b[2]]
            ncol = len(headers)
            pdf.set_x(18); pdf.set_font('Ar', '', 9.5)
            pdf.set_fill_color(255, 255, 255)  # évite la « bavure » d'un encadré précédent
            head_style = FontFace(emphasis="BOLD", color=(255, 255, 255), fill_color=NAVY)
            kw = dict(width=W, headings_style=head_style, line_height=5.2, padding=1.6, markdown=True,
                      cell_fill_color=(244, 246, 249), cell_fill_mode="ROWS", borders_layout="MINIMAL")
            if ncol == 3:
                kw['col_widths'] = (16, 40, 44); kw['text_align'] = ("CENTER", "LEFT", "LEFT")
            else:
                kw['text_align'] = "LEFT"
            with pdf.table(**kw) as table:
                r = table.row()
                for h in headers:
                    r.cell(h)
                for data in rows:
                    r = table.row()
                    for c in data:
                        r.cell(c)
            pdf.ln(4)
        elif kind == 'takeaways':
            pts = [pick(pt, lang) for pt in b[1]]
            th = mc_height(pdf, L['takeaways'].upper(), W - 3, 7, 'B', 11) + 4
            for pt in pts:
                th += mc_height(pdf, "•  " + pt, W - 3, 5.6, '', 10, markdown=True)
            if pdf.will_page_break(th):
                pdf.add_page()
            pdf.ln(2); x0 = 18; y0 = pdf.get_y()
            pdf.set_fill_color(250, 245, 230); pdf.set_x(x0 + 3)
            pdf.set_font('Ar', 'B', 11); pdf.set_text_color(*GOLD_D)
            pdf.multi_cell(W - 3, 7, L['takeaways'].upper(), fill=True)
            pdf.set_font('Ar', '', 10); pdf.set_text_color(*INK)
            for pt in pts:
                pdf.set_x(x0 + 3); pdf.multi_cell(W - 3, 5.6, "•  " + pt, fill=True, markdown=True)
            y1 = pdf.get_y(); pdf.set_draw_color(*GOLD); pdf.set_line_width(0.6)
            pdf.rect(x0, y0, W, y1 - y0); pdf.ln(2)
    pdf.output(path)


# ======================= DOCX (python-docx) =======================
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


def add_runs(par, text, size=11, color=INK, italic=False):
    for seg, bold in split_bold(text):
        r = par.add_run(seg); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(*color); r.font.name = 'Calibri'


def box(doc, hexfill):
    t = doc.add_table(rows=1, cols=1); t.autofit = True
    cell = t.cell(0, 0); shade(cell, hexfill)
    cell.paragraphs[0].text = ''
    return cell


def render_docx(num, ch, lang, path):
    L = LABELS[lang]
    title = pick(ch['title'], lang)
    subtitle = pick(ch.get('subtitle'), lang) or L['subtitle']

    doc = docx.Document()
    sec = doc.sections[0]
    sec.top_margin = Mm(16); sec.bottom_margin = Mm(16); sec.left_margin = Mm(18); sec.right_margin = Mm(18)

    # Titre
    p = doc.add_paragraph(); r = p.add_run(L['kicker'] % num)
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = RGBColor(*GOLD_D); r.font.name = 'Calibri'
    p = doc.add_paragraph(); r = p.add_run(title)
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(*NAVY); r.font.name = 'Calibri'
    p = doc.add_paragraph(); r = p.add_run(subtitle)
    r.font.size = Pt(11); r.font.color.rgb = RGBColor(*SLATE); r.italic = True; r.font.name = 'Calibri'

    # Sommaire
    cell = box(doc, 'F4F6F9')
    pp = cell.paragraphs[0]; rr = pp.add_run(L['toc'])
    rr.bold = True; rr.font.size = Pt(11); rr.font.color.rgb = RGBColor(*NAVY); rr.font.name = 'Calibri'
    for t in toc(ch['blocks'], lang):
        sp = cell.add_paragraph(); r = sp.add_run("•  " + t)
        r.font.size = Pt(10); r.font.color.rgb = RGBColor(*SLATE); r.font.name = 'Calibri'
    doc.add_paragraph()

    # Blocs
    for b in ch['blocks']:
        kind = b[0]
        if kind == 'h2':
            heading = pick(b[1], lang)
            p = doc.add_paragraph(style='Heading 1')  # vrai style Word -> volet de navigation + accessibilité
            r = p.add_run(heading.replace('**', ''))
            r.bold = True; r.font.size = Pt(14); r.font.color.rgb = RGBColor(*NAVY); r.font.name = 'Calibri'
            p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
            # filet doré
            pPr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '2'); bottom.set(qn('w:color'), 'C2A14D')
            pbdr.append(bottom); pPr.append(pbdr)
        elif kind == 'p':
            add_runs(doc.add_paragraph(), pick(b[1], lang))
        elif kind == 'callout':
            _, ctype, ctitle, ctext = b
            ctitle = pick(ctitle, lang); ctext = pick(ctext, lang)
            bd, _bg, hx = CALLOUT.get(ctype, CALLOUT['info'])
            cell = box(doc, hx)
            tp = cell.paragraphs[0]; r = tp.add_run(ctitle.replace('**', ''))
            r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(*bd); r.font.name = 'Calibri'
            add_runs(cell.add_paragraph(), ctext, size=10)
            doc.add_paragraph()
        elif kind == 'lawref':
            _, ref, txt = b
            ref = pick(ref, lang); txt = pick(txt, lang)
            bd, _bg, hx = CALLOUT['legal']
            cell = box(doc, hx)
            tp = cell.paragraphs[0]
            r = tp.add_run((L['lawref'] + " — " + ref_label(ref)).replace('**', ''))
            r.bold = True; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(*bd); r.font.name = 'Calibri'
            add_runs(cell.add_paragraph(), "« " + txt + " »", size=10, color=SLATE, italic=True)
            doc.add_paragraph()
        elif kind == 'table':
            headers = [pick(h, lang) for h in b[1]]
            rows = [[pick(c, lang) for c in row] for row in b[2]]
            t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'
            for i, h in enumerate(headers):
                c = t.rows[0].cells[i]; shade(c, '1B3A5B')
                rn = c.paragraphs[0].add_run(h.replace('**', ''))
                rn.bold = True; rn.font.size = Pt(9.5); rn.font.color.rgb = RGBColor(255, 255, 255); rn.font.name = 'Calibri'
            for ri, data in enumerate(rows):
                cells = t.add_row().cells
                for i, val in enumerate(data):
                    if ri % 2:
                        shade(cells[i], 'F4F6F9')
                    add_runs(cells[i].paragraphs[0], val, size=9.5)
            doc.add_paragraph()
        elif kind == 'takeaways':
            cell = box(doc, 'FAF5E6')
            tp = cell.paragraphs[0]; r = tp.add_run(L['takeaways'])
            r.bold = True; r.font.size = Pt(12); r.font.color.rgb = RGBColor(*GOLD_D); r.font.name = 'Calibri'
            for pt in b[1]:
                add_runs(cell.add_paragraph(), "•  " + pick(pt, lang), size=10)

    # Pied de page
    ftp = sec.footer.paragraphs[0]; ftp.text = L['footer']
    for r in ftp.runs:
        r.font.size = Pt(8); r.font.color.rgb = RGBColor(*SLATE)
    doc.save(path)


# ======================= Orchestration =======================
def main(argv):
    arg = (argv[1].lower() if len(argv) > 1 else 'both')
    if arg in ('fr', 'en'):
        langs = (arg,)
    elif arg == 'both':
        langs = LANGS
    else:
        raise SystemExit("Langue inconnue : %r (attendu : 'fr', 'en' ou 'both')" % arg)

    os.makedirs(OUT, exist_ok=True)
    chapters = load_chapters()

    for lang in langs:
        for num, ch in chapters.items():
            pdfp = os.path.join(OUT, "chapitre-%d-polycopie.%s.pdf" % (num, lang))
            docp = os.path.join(OUT, "chapitre-%d-polycopie.%s.docx" % (num, lang))
            render_pdf(num, ch, lang, pdfp)
            render_docx(num, ch, lang, docp)
            print("OK [%s] chapitre %d : %d o (pdf) %d o (docx)"
                  % (lang, num, os.path.getsize(pdfp), os.path.getsize(docp)))


if __name__ == '__main__':
    main(sys.argv)
