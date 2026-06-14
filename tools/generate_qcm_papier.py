# -*- coding: utf-8 -*-
"""Template de site de cours — Mohamed EL AFRIT (MIT / CC BY 4.0)

Genere la version PAPIER du QCM a partir des donnees BILINGUES du site
(site/QCM/questions.js + corrections.js), pour la ou les langues demandees.

Pour chaque langue (fr / en), produit dans tools/out/ :
  - QCM_questions.<lang>.docx       : page de garde + consignes + bareme
                                      (avec exemples chiffres) + questions (2 colonnes)
  - QCM_grille_reponses.<lang>.docx : grille de reponses (N° + A/B/C/D + certitude F/M/É)
  - QCM_corrige.<lang>.docx         : corrige formateur (bonne reponse + explication)

Donnees bilingues attendues :
  meta.titre={fr,en}, question.enonce={fr,en}, question.propositions=[{fr,en}...],
  correction.explication={fr,en}, correction.bonne_reponse_index.

Bareme CBM : Élevée +3/−3 · Moyenne +2/−2 · Faible +1/0.
Note /20 = score brut ÷ (3×N) × 20.

Usage :
  python generate_qcm_papier.py [fr|en|both]
  (defaut : both)
"""
import os, sys, json, re

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------- chemins ----------------
HERE = os.path.dirname(os.path.abspath(__file__))                 # .../tools
ROOT = os.path.dirname(HERE)                                      # .../template_cours
SITE_QCM = os.path.join(ROOT, 'site', 'QCM')
OUT_DIR = os.path.join(HERE, 'out')

QUESTIONS_JS = os.path.join(SITE_QCM, 'questions.js')
CORRECTIONS_JS = os.path.join(SITE_QCM, 'corrections.js')

L = ['A', 'B', 'C', 'D']
GRAY = 'D9D9D9'
LGRAY = 'F2F2F2'

# ---------------- extraction des donnees bilingues ----------------
def _extract_js_object(path, marker):
    """Lit un fichier .js de la forme `window.X = { ... };` et renvoie le dict.
    On decoupe sur la 1re { apres le marqueur et la derniere } avant le ; final.
    Le JS de ce site utilise des cles non quotees et des apostrophes simples :
    on normalise donc le fragment en JSON avant json.loads.
    """
    with open(path, encoding='utf-8') as f:
        src = f.read()
    pos = src.find(marker)
    if pos == -1:
        raise ValueError("Marqueur introuvable dans %s : %r" % (path, marker))
    frag = src[pos + len(marker):]
    start = frag.find('{')
    end = frag.rfind('}')
    if start == -1 or end == -1 or end <= start:
        raise ValueError("Bloc { ... } introuvable dans %s" % path)
    obj = frag[start:end + 1]
    return _js_object_to_dict(obj)


def _js_object_to_dict(obj):
    """Convertit un litteral objet JS (cles non quotees, apostrophes simples,
    commentaires /* */ et //, virgules finales) en dict Python."""
    s = obj
    # 1) retirer les commentaires /* ... */ et // ...
    s = re.sub(r'/\*.*?\*/', '', s, flags=re.S)
    s = re.sub(r'//[^\n\r]*', '', s)

    # 2) tokeniser en respectant les chaines (simples ou doubles) pour
    #    transformer les chaines 'single' en "double" et quoter les cles.
    out = []
    i = 0
    n = len(s)
    while i < n:
        ch = s[i]
        if ch in ("'", '"'):
            quote = ch
            j = i + 1
            buf = []
            while j < n:
                c = s[j]
                if c == '\\' and j + 1 < n:
                    nxt = s[j + 1]
                    # garder l'echappement utile ; \' devient ' en sortie double
                    if nxt == "'":
                        buf.append("'")
                    elif nxt == '"':
                        buf.append('\\"')
                    elif nxt == '\\':
                        buf.append('\\\\')
                    elif nxt in ('n', 't', 'r', '/', 'b', 'f', 'u'):
                        buf.append('\\' + nxt)
                    else:
                        buf.append(nxt)
                    j += 2
                    continue
                if c == quote:
                    j += 1
                    break
                if c == '"':
                    buf.append('\\"')
                elif c == '\n':
                    buf.append('\\n')
                elif c == '\r':
                    pass
                elif c == '\t':
                    buf.append('\\t')
                else:
                    buf.append(c)
                j += 1
            out.append('"' + ''.join(buf) + '"')
            i = j
        else:
            out.append(ch)
            i += 1
    s = ''.join(out)

    # 3) quoter les cles d'objet non quotees : { cle: ...  ou  , cle: ...
    s = re.sub(r'([{,]\s*)([A-Za-z_$][\w$]*)(\s*):', r'\1"\2"\3:', s)

    # 4) supprimer les virgules finales avant } ou ]
    s = re.sub(r',(\s*[}\]])', r'\1', s)

    return json.loads(s)


def load_data():
    q = _extract_js_object(QUESTIONS_JS, 'window.QCM_QUESTIONS')
    c = _extract_js_object(CORRECTIONS_JS, 'window.QCM_CORRECTIONS')
    meta = q.get('meta', {})
    questions = q.get('questions', [])
    corrections = c.get('corrections', [])
    # index des corrections par id (sinon par position)
    by_id = {}
    for k, corr in enumerate(corrections):
        cid = corr.get('id')
        if cid is not None:
            by_id[cid] = corr
    merged = []
    for k, item in enumerate(questions):
        qid = item.get('id')
        corr = by_id.get(qid)
        if corr is None and k < len(corrections):
            corr = corrections[k]
        if corr is None:
            corr = {}
        merged.append((item, corr))
    return meta, merged


# ---------------- libelles bilingues (consignes / barème / textes) ----------------
def labels(lang, titre, sous, duree, n, maxraw, auteur):
    if lang == 'en':
        return {
            'cover_kicker': "QUIZ — INDIVIDUAL ASSESSMENT",
            'cover_meta': sous + "  ·  Time: " + duree + "  ·  " + str(n) + " questions",
            'idline': "LAST NAME: ______________________________     FIRST NAME: ______________________________     DATE: ____ / ____ / ______",
            'instructions_h': "Instructions",
            'instructions': [
                "Individual assessment (duration " + duree + "). No resources allowed: neither the Internet nor any artificial-intelligence tool; only a non-connected calculator and a double-sided sheet of notes are permitted.",
                "The quiz has " + str(n) + " single-answer questions: for each question, exactly one option (A, B, C or D) is correct.",
                "Report all your answers on the answer sheet, the only document collected. For each question, fill in a single answer circle (A/B/C/D) and a single certainty circle (L/M/H).",
                "Every mark must be clear; in case of error, cross it out clearly. If no certainty is indicated, the “Low” certainty is applied.",
            ],
            'scale_h': "Marking scheme — certainty-based marking (CBM)",
            'scale_intro': "You state your certainty for each answer. The more confident you are, the more a correct answer earns… but the more a mistake costs.",
            'scale_hdr': ["Your certainty", "Correct answer", "Wrong answer"],
            'scale_rows': [
                ["High  (you are certain)", "+3", "−3"],
                ["Medium  (fairly sure)", "+2", "−2"],
                ["Low  (you are guessing)", "+1", "0"],
            ],
            'note_formula_b': "Mark /20 = raw score ÷ " + str(maxraw) + " × 20",
            'note_formula_rest': "   (raw score = sum of points; maximum " + str(maxraw) + " for " + str(n) + " correct answers at High certainty; a negative mark is raised to 0).",
            'examples_h': "Examples",
            'examples': [
                "Correct answer at High certainty → +3. The same answer wrong → −3.",
                "Correct answer at Medium → +2; wrong → −2.   Correct answer at Low → +1; wrong → 0 (no penalty when you own up to a guess).",
            ],
            'example_calc_intro': "Worked example on 3 questions:",
            'example_table': [
                ["Question", "Your answer", "Certainty", "Points"],
                ["1", "correct", "High", "+3"],
                ["2", "incorrect", "Medium", "−2"],
                ["3", "correct", "Low", "+1"],
                ["", "", "Total (3 q.)", "+2"],
            ],
            'example_global': "Across all " + str(n) + " questions, a raw score of " + _example_raw(n, maxraw) + " gives, for instance: " + _example_raw(n, maxraw) + " ÷ " + str(maxraw) + " × 20 = " + _example_mark(n, maxraw) + " / 20.",
            'strategy': "Honest strategy: “Low” if you are guessing, “Medium” if you are fairly sure, “High” only if you are truly certain — a wrong High certainty costs −3.",
            'questions_h': "Questions (1 to " + str(n) + ")",
            'foot_q': titre + " — Quiz — " + auteur + " — CC BY 4.0",
            # grille
            'grille_title2': "Quiz — Answer sheet   (the only document to hand in)",
            'grille_instructions': [
                ("Single answer: fill in a single circle A, B, C or D per question.", False),
                ("Certainty (one circle ○): L = Low · M = Medium · H = High   (default: Low).", False),
                ("Scale: High +3 / −3  ·  Medium +2 / −2  ·  Low +1 / 0.   Mark /20 = raw ÷ " + str(maxraw) + " × 20.", True),
            ],
            'grille_sub': ['No.', 'A', 'B', 'C', 'D', 'L', 'M', 'H'],
            'foot_grille': titre + " — Answer sheet — " + auteur + " — CC BY 4.0",
            # corrige
            'corrige_kicker': "ANSWER KEY — for the instructor only",
            'corrige_intro': "Quiz — " + str(n) + " single-answer questions. CBM scale: High +3/−3 · Medium +2/−2 · Low +1/0. Mark /20 = raw ÷ " + str(maxraw) + " × 20.",
            'corrige_chapter': "Chapter ",
            'corrige_answer': "Answer: ",
            'corrige_expl': "Explanation. ",
            'foot_corrige': titre + " — Answer key (instructor) — " + auteur + " — CC BY 4.0",
        }
    # ----- FR (defaut) -----
    return {
        'cover_kicker': "QCM — ÉPREUVE INDIVIDUELLE",
        'cover_meta': sous + "  ·  Durée : " + duree + "  ·  " + str(n) + " questions",
        'idline': "NOM : ______________________________     PRÉNOM : ______________________________     DATE : ____ / ____ / ______",
        'instructions_h': "Consignes",
        'instructions': [
            "Épreuve individuelle (durée " + duree + "). Aucune ressource autorisée : ni Internet, ni outil d'intelligence artificielle ; seules une calculatrice non connectée et une feuille de notes recto-verso sont admises.",
            "Le questionnaire comporte " + str(n) + " questions à réponse unique : pour chaque question, une seule proposition (A, B, C ou D) est correcte.",
            "Reportez toutes vos réponses sur la grille de réponses, seul document ramassé. Pour chaque question, noircissez un seul cercle de réponse (A/B/C/D) et un seul cercle de certitude (F/M/É).",
            "Toute marque doit être nette ; en cas d'erreur, raturez clairement. À défaut de certitude indiquée, la certitude « Faible » est appliquée.",
        ],
        'scale_h': "Barème — notation par degré de certitude (CBM)",
        'scale_intro': "Vous indiquez votre certitude pour chaque réponse. Plus vous êtes sûr, plus une bonne réponse rapporte… mais plus une erreur coûte.",
        'scale_hdr': ["Votre certitude", "Bonne réponse", "Mauvaise réponse"],
        'scale_rows': [
            ["Élevée  (vous êtes certain)", "+3", "−3"],
            ["Moyenne  (assez sûr)", "+2", "−2"],
            ["Faible  (vous devinez)", "+1", "0"],
        ],
        'note_formula_b': "Note /20 = score brut ÷ " + str(maxraw) + " × 20",
        'note_formula_rest': "   (score brut = somme des points ; maximum " + str(maxraw) + " pour " + str(n) + " bonnes réponses en certitude Élevée ; une note négative est ramenée à 0).",
        'examples_h': "Exemples",
        'examples': [
            "Bonne réponse en certitude Élevée → +3. La même réponse fausse → −3.",
            "Bonne réponse en Moyenne → +2 ; fausse → −2.   Bonne réponse en Faible → +1 ; fausse → 0 (aucune pénalité quand on assume un pari).",
        ],
        'example_calc_intro': "Exemple de calcul sur 3 questions :",
        'example_table': [
            ["Question", "Votre réponse", "Certitude", "Points"],
            ["1", "correcte", "Élevée", "+3"],
            ["2", "incorrecte", "Moyenne", "−2"],
            ["3", "correcte", "Faible", "+1"],
            ["", "", "Total (3 q.)", "+2"],
        ],
        'example_global': "Sur l'ensemble des " + str(n) + " questions, un score brut de " + _example_raw(n, maxraw) + " donne par exemple : " + _example_raw(n, maxraw) + " ÷ " + str(maxraw) + " × 20 = " + _example_mark(n, maxraw) + " / 20.",
        'strategy': "Stratégie honnête : « Faible » si vous devinez, « Moyenne » si vous êtes assez sûr, « Élevée » seulement si vous êtes vraiment certain — une certitude Élevée erronée coûte −3.",
        'questions_h': "Questions (1 à " + str(n) + ")",
        'foot_q': titre + " — QCM — " + auteur + " — CC BY 4.0",
        # grille
        'grille_title2': "QCM — Grille de réponses   (seul document à rendre)",
        'grille_instructions': [
            ("Réponse unique : noircissez un seul cercle A, B, C ou D par question.", False),
            ("Certitude (un seul cercle ○) : F = Faible · M = Moyenne · É = Élevée   (à défaut : Faible).", False),
            ("Barème : Élevée +3 / −3  ·  Moyenne +2 / −2  ·  Faible +1 / 0.   Note /20 = brut ÷ " + str(maxraw) + " × 20.", True),
        ],
        'grille_sub': ['N°', 'A', 'B', 'C', 'D', 'F', 'M', 'É'],
        'foot_grille': titre + " — Grille QCM — " + auteur + " — CC BY 4.0",
        # corrige
        'corrige_kicker': "CORRIGÉ — réservé au formateur",
        'corrige_intro': "QCM — " + str(n) + " questions à réponse unique. Barème CBM : Élevée +3/−3 · Moyenne +2/−2 · Faible +1/0. Note /20 = brut ÷ " + str(maxraw) + " × 20.",
        'corrige_chapter': "Chapitre ",
        'corrige_answer': "Réponse : ",
        'corrige_expl': "Explication. ",
        'foot_corrige': titre + " — Corrigé QCM (formateur) — " + auteur + " — CC BY 4.0",
    }


def _example_raw(n, maxraw):
    # un score brut "credible" : ~70% du maximum
    return str(int(round(maxraw * 0.7)))


def _example_mark(n, maxraw):
    raw = int(round(maxraw * 0.7))
    val = raw / maxraw * 20.0
    s = ("%.1f" % val).replace('.', ',')
    return s


# ---------------- helpers de mise en page (logique du script source) ----------------
def set_margins(sec, t=13, b=12, l=13, r=13):
    sec.top_margin = Mm(t); sec.bottom_margin = Mm(b); sec.left_margin = Mm(l); sec.right_margin = Mm(r)


def set_cols(sec, num, space=420):
    sectPr = sec._sectPr
    cols = sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols'); sectPr.append(cols)
    cols.set(qn('w:num'), str(num)); cols.set(qn('w:space'), str(space))


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def cell_w(cell, tw):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(tw)); w.set(qn('w:type'), 'dxa'); tcPr.append(w)


def fixed_layout(table, widths):
    table.autofit = False; table.allow_autofit = False
    tblPr = table._tbl.tblPr
    lay = tblPr.find(qn('w:tblLayout'))
    if lay is None:
        lay = OxmlElement('w:tblLayout'); tblPr.append(lay)
    lay.set(qn('w:type'), 'fixed')
    grid = table._tbl.find(qn('w:tblGrid'))
    gcols = grid.findall(qn('w:gridCol'))
    for gc, w in zip(gcols, widths): gc.set(qn('w:w'), str(w))
    for row in table.rows:
        for c, w in zip(row.cells, widths): cell_w(c, w)


def no_space(par, before=0, after=0, line=1.0):
    pf = par.paragraph_format; pf.space_before = Pt(before); pf.space_after = Pt(after)
    pf.line_spacing = line


def run(par, text, size=9, bold=False, italic=False, color=(0, 0, 0), font='Calibri'):
    r = par.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.name = font; r.font.color.rgb = RGBColor(*color)
    return r


def center_cell(cell):
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = 1  # center


def tight_cell(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for side in ('top', 'bottom', 'start', 'end'):
        e = OxmlElement('w:' + side); e.set(qn('w:w'), '15'); e.set(qn('w:type'), 'dxa'); m.append(e)
    tcPr.append(m)


def pick(value, lang):
    """Renvoie la variante de langue d'un champ {fr,en} ; tolere une chaine simple."""
    if isinstance(value, dict):
        return value.get(lang) or value.get('fr') or value.get('en') or ''
    return value if value is not None else ''


# ======================= 1) QCM (questions) =======================
def build_questions(meta, merged, lang, ctx):
    titre, sous, duree, auteur, n, maxraw = ctx['titre'], ctx['sous'], ctx['duree'], ctx['auteur'], ctx['n'], ctx['maxraw']
    T = labels(lang, titre, sous, duree, n, maxraw, auteur)
    doc = Document()
    s0 = doc.sections[0]; set_margins(s0)
    # --- Page de garde ---
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p, 0, 1)
    run(p, T['cover_kicker'], 11, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p, 0, 1)
    run(p, titre, 16, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p, 0, 6)
    run(p, T['cover_meta'], 10)
    # NOM / PRENOM
    p = doc.add_paragraph(); no_space(p, 2, 6)
    run(p, T['idline'], 10)

    # --- Consignes ---
    p = doc.add_paragraph(); no_space(p, 4, 2); run(p, T['instructions_h'], 12, bold=True)
    for txt in T['instructions']:
        q = doc.add_paragraph(); no_space(q, 0, 2); run(q, "•  " + txt, 9.5)

    # --- Bareme ---
    p = doc.add_paragraph(); no_space(p, 6, 2); run(p, T['scale_h'], 12, bold=True)
    q = doc.add_paragraph(); no_space(q, 0, 3)
    run(q, T['scale_intro'], 9.5)
    t = doc.add_table(rows=4, cols=3); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = T['scale_hdr']
    rows = T['scale_rows']
    for j, h in enumerate(hdr):
        c = t.rows[0].cells[j]; shade(c, GRAY); tight_cell(c)
        c.paragraphs[0].text = ''; run(c.paragraphs[0], h, 9, bold=True)
    for i, rdata in enumerate(rows):
        for j, val in enumerate(rdata):
            c = t.rows[i + 1].cells[j]; tight_cell(c)
            c.paragraphs[0].text = ''; run(c.paragraphs[0], val, 9, bold=(j == 0))
            if j > 0: c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for col, w in zip(t.columns, (4400, 2400, 2600)):
        for c in col.cells: cell_w(c, w)
    fixed_layout(t, (4400, 2400, 2600))

    q = doc.add_paragraph(); no_space(q, 4, 2)
    run(q, T['note_formula_b'], 10, bold=True)
    run(q, T['note_formula_rest'], 9.5)

    # exemples
    p = doc.add_paragraph(); no_space(p, 4, 2); run(p, T['examples_h'], 10.5, bold=True)
    for txt in T['examples']:
        q = doc.add_paragraph(); no_space(q, 0, 1); run(q, "•  " + txt, 9.5)
    # mini tableau exemple chiffre
    q = doc.add_paragraph(); no_space(q, 2, 2); run(q, T['example_calc_intro'], 9.5, italic=True)
    ex = T['example_table']
    te = doc.add_table(rows=len(ex), cols=4); te.style = 'Table Grid'
    last = len(ex) - 1
    for i, rdata in enumerate(ex):
        for j, val in enumerate(rdata):
            c = te.rows[i].cells[j]; tight_cell(c)
            c.paragraphs[0].text = ''
            run(c.paragraphs[0], val, 8.5, bold=(i == 0 or i == last))
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if i == 0: shade(c, GRAY)
            if i == last and j >= 2: shade(c, LGRAY)
    fixed_layout(te, (1700, 2700, 2700, 1700))
    q = doc.add_paragraph(); no_space(q, 3, 2); run(q, T['example_global'], 9.5, italic=True)
    q = doc.add_paragraph(); no_space(q, 2, 0)
    run(q, T['strategy'], 9.5, bold=True)

    # --- Section 2 colonnes : les questions ---
    doc.add_section(WD_SECTION.NEW_PAGE)
    s1 = doc.sections[-1]; set_margins(s1); set_cols(s1, 2, 420)
    head = doc.add_paragraph(); no_space(head, 0, 4); run(head, T['questions_h'], 11, bold=True)
    for idx, (item, corr) in enumerate(merged):
        num = idx + 1
        pe = doc.add_paragraph(); no_space(pe, 5, 1, 1.0)
        pe.paragraph_format.keep_with_next = True
        run(pe, str(num) + ". ", 9, bold=True); run(pe, pick(item.get('enonce'), lang), 9)
        props = item.get('propositions', [])
        for i, prop in enumerate(props):
            pp = doc.add_paragraph(); no_space(pp, 0, 0, 1.0)
            pp.paragraph_format.left_indent = Mm(4)
            if i < len(props) - 1: pp.paragraph_format.keep_with_next = True
            lab = L[i] if i < len(L) else chr(ord('A') + i)
            run(pp, lab + ". ", 8.5, bold=True); run(pp, pick(prop, lang), 8.5)
    # pied
    f = doc.sections[-1].footer.paragraphs[0]; f.text = ''
    run(f, T['foot_q'], 7.5, color=(90, 90, 90))
    out = os.path.join(OUT_DIR, 'QCM_questions.%s.docx' % lang)
    doc.save(out)
    return out


# ======================= 2) Grille de reponses =======================
def _balanced_blocks(n, max_blocks=3):
    """Repartit n questions en <= max_blocks colonnes equilibrees.
    Renvoie une liste de (premier_index_0based, nombre) par bloc."""
    blocks = min(max_blocks, n) if n > 0 else 1
    per = (n + blocks - 1) // blocks  # plafond
    blocks = (n + per - 1) // per     # recalcule le nb de blocs reellement utiles
    res = []
    start = 0
    for b in range(blocks):
        cnt = min(per, n - start)
        if cnt <= 0:
            break
        res.append((start, cnt))
        start += cnt
    return res, per


def build_grille(meta, merged, lang, ctx):
    titre, sous, auteur, n, maxraw = ctx['titre'], ctx['sous'], ctx['auteur'], ctx['n'], ctx['maxraw']
    T = labels(lang, titre, ctx['sous'], ctx['duree'], n, maxraw, auteur)
    doc = Document(); s = doc.sections[0]; set_margins(s, 12, 10, 12, 12)
    # bandeau titre
    bt = doc.add_table(rows=1, cols=1); bt.style = 'Table Grid'
    c = bt.rows[0].cells[0]; shade(c, GRAY); tight_cell(c); c.paragraphs[0].text = ''
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(c.paragraphs[0], titre + "  ·  " + sous, 11, bold=True)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p2, 0, 0)
    run(p2, T['grille_title2'], 9.5)
    # NOM
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    pn = doc.add_paragraph(); no_space(pn, 0, 4)
    run(pn, T['idline'], 10)
    # consignes condensees
    for txt, bold in T['grille_instructions']:
        q = doc.add_paragraph(); no_space(q, 0, 1); run(q, txt, 9, bold=bold)

    # grille : blocs equilibres ; chaque bloc = [N° | A B C D | F M É]
    blocks, per = _balanced_blocks(n, 3)
    nb = len(blocks)
    sub = T['grille_sub']
    cols = nb * len(sub)
    rows = per + 1  # header + lignes
    t = doc.add_table(rows=rows, cols=cols); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # en-tete
    for b in range(nb):
        for k, lab in enumerate(sub):
            c = t.rows[0].cells[b * len(sub) + k]; tight_cell(c); center_cell(c)
            shade(c, GRAY if k <= 4 else LGRAY)
            c.paragraphs[0].text = ''; run(c.paragraphs[0], lab, 8.5, bold=True)
    # corps
    for r in range(per):
        for b in range(nb):
            start, cnt = blocks[b]
            base = b * len(sub)
            if r >= cnt:
                # cellules vides (bloc plus court) : on laisse vide mais tracé
                for k in range(len(sub)):
                    c = t.rows[r + 1].cells[base + k]; tight_cell(c); center_cell(c)
                    if k >= 5: shade(c, LGRAY)
                    c.paragraphs[0].text = ''
                continue
            qnum = start + r + 1
            cN = t.rows[r + 1].cells[base]; tight_cell(cN); center_cell(cN)
            cN.paragraphs[0].text = ''; run(cN.paragraphs[0], str(qnum), 8.5, bold=True)
            for k in range(1, 8):
                c = t.rows[r + 1].cells[base + k]; tight_cell(c); center_cell(c)
                if k >= 5: shade(c, LGRAY)
                c.paragraphs[0].text = ''; run(c.paragraphs[0], '○', 9)
    # largeurs : N°=520, lettres=300
    widths = []
    for b in range(nb): widths += [520] + [300] * 7
    fixed_layout(t, widths)
    # pied
    f = doc.sections[0].footer.paragraphs[0]; f.text = ''
    run(f, T['foot_grille'], 7.5, color=(90, 90, 90))
    out = os.path.join(OUT_DIR, 'QCM_grille_reponses.%s.docx' % lang)
    doc.save(out)
    return out


# ======================= 3) Corrige formateur =======================
def build_corrige(meta, merged, lang, ctx):
    titre, sous, auteur, n, maxraw = ctx['titre'], ctx['sous'], ctx['auteur'], ctx['n'], ctx['maxraw']
    T = labels(lang, titre, ctx['sous'], ctx['duree'], n, maxraw, auteur)
    doc = Document(); s = doc.sections[0]; set_margins(s, 14, 12, 14, 14)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p, 0, 1)
    run(p, T['corrige_kicker'], 11, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; no_space(p, 0, 2)
    run(p, titre + "  ·  " + sous, 12, bold=True)
    q = doc.add_paragraph(); no_space(q, 0, 6)
    run(q, T['corrige_intro'], 9.5, italic=True)
    cur_ch = None
    for idx, (item, corr) in enumerate(merged):
        num = idx + 1
        ch = item.get('chapitre', corr.get('chapitre'))
        if ch is not None and ch != cur_ch:
            cur_ch = ch
            h = doc.add_paragraph(); no_space(h, 6, 2)
            run(h, T['corrige_chapter'] + str(cur_ch), 11, bold=True)
        pe = doc.add_paragraph(); no_space(pe, 3, 1)
        run(pe, str(num) + ". ", 9.5, bold=True); run(pe, pick(item.get('enonce'), lang), 9.5)
        bi = corr.get('bonne_reponse_index')
        props = item.get('propositions', [])
        pr = doc.add_paragraph(); no_space(pr, 0, 1); pr.paragraph_format.left_indent = Mm(4)
        if bi is not None and 0 <= bi < len(props):
            lab = L[bi] if bi < len(L) else chr(ord('A') + bi)
            run(pr, T['corrige_answer'] + lab + " — ", 9.5, bold=True)
            run(pr, pick(props[bi], lang), 9.5, bold=True)
        else:
            run(pr, T['corrige_answer'] + "?", 9.5, bold=True)
        expl = pick(corr.get('explication'), lang)
        if expl:
            px = doc.add_paragraph(); no_space(px, 0, 1); px.paragraph_format.left_indent = Mm(4)
            run(px, T['corrige_expl'], 9, italic=True, color=(80, 80, 80))
            run(px, expl, 9, color=(40, 40, 40))
    f = doc.sections[0].footer.paragraphs[0]; f.text = ''
    run(f, T['foot_corrige'], 7.5, color=(90, 90, 90))
    out = os.path.join(OUT_DIR, 'QCM_corrige.%s.docx' % lang)
    doc.save(out)
    return out


# ---------------- orchestration ----------------
def build_context(meta, merged, lang):
    n = len(merged)
    maxraw = 3 * n if n else 0
    titre = pick(meta.get('titre'), lang) or ("QCM" if lang == 'fr' else "Quiz")
    formation = pick(meta.get('formation'), lang)
    annee = meta.get('annee_academique', '')
    sep = "  ·  "
    sous = sep.join([x for x in (formation, annee) if x])
    if not sous:
        sous = annee or formation or ""
    duree = meta.get('duree')
    if not duree:
        duree = "1 h 30" if lang == 'fr' else "1 h 30"
    auteur = meta.get('auteur', 'Mohamed EL AFRIT')
    return {'titre': titre, 'sous': sous, 'duree': duree, 'auteur': auteur, 'n': n, 'maxraw': maxraw}


def run_lang(meta, merged, lang):
    ctx = build_context(meta, merged, lang)
    outs = [
        build_questions(meta, merged, lang, ctx),
        build_grille(meta, merged, lang, ctx),
        build_corrige(meta, merged, lang, ctx),
    ]
    return outs


def main():
    arg = (sys.argv[1].lower() if len(sys.argv) > 1 else 'both')
    if arg in ('fr', 'en'):
        langs = [arg]
    elif arg in ('both', 'all', 'fr+en'):
        langs = ['fr', 'en']
    else:
        print("Argument de langue inconnu : %r (attendu : fr | en | both). Defaut : both." % arg)
        langs = ['fr', 'en']

    os.makedirs(OUT_DIR, exist_ok=True)
    meta, merged = load_data()
    if not merged:
        print("Aucune question chargee depuis", QUESTIONS_JS)
        return

    for lang in langs:
        outs = run_lang(meta, merged, lang)
        for f in outs:
            try:
                size = os.path.getsize(f)
            except OSError:
                size = 0
            print('OK', f, size, 'o')


if __name__ == '__main__':
    main()
